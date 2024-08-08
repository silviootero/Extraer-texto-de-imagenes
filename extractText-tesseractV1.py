
import os
from PIL import Image
import pytesseract
from docx import Document

# Especificar la ruta completa al ejecutable de Tesseract
pytesseract.pytesseract.tesseract_cmd = r'D:/TESSERACT/tesseract.exe'

# Ruta a la carpeta que contiene las imágenes
folder_path = 'C:/Users/hp/Desktop/CAPS_TRADUS'

# Crear un nuevo documento de Word
doc = Document()

# Función para ordenar archivos numéricamente
def sort_numerically(value):
    parts = value.split('.')
    num = ''.join(filter(str.isdigit, parts[0]))
    return int(num)

# Obtener y ordenar los archivos numéricamente
files = [f for f in os.listdir(folder_path) if f.endswith(('.png', '.jpg', '.jpeg'))]
files.sort(key=sort_numerically)

# Recorrer cada archivo en la carpeta ordenada
for filename in files:
    # Ruta completa al archivo de imagen
    image_path = os.path.join(folder_path, filename)

    # Abrir imagen
    image = Image.open(image_path)

    # Usar pytesseract para extraer texto
    text = pytesseract.image_to_string(image)

    # Añadir el texto al documento
    doc.add_heading(filename, level=1)  # Añadir el nombre del archivo como encabezado
    doc.add_paragraph(text)  # Añadir el texto extraído

# Guardar el documento
doc.save('TEXTO.docx')
