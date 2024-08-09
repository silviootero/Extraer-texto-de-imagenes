
import os
import cv2
import numpy as np
from PIL import Image
from docx import Document
import easyocr


# Ruta a la carpeta que contiene las imágenes
folder_path = 'C:/Users/hp/Desktop/CAPS_TRADUS'

# Crear un nuevo documento de Word
doc = Document()

# Función para ordenar archivos numéricamente
def sort_numerically(value):
    parts = value.split('.')
    num = ''.join(filter(str.isdigit, parts[0]))
    return int(num)

# Obtener y ordenar los archivos numéricamente
files = [f for f in os.listdir(folder_path) if f.endswith(('.png', '.jpg', '.jpeg'))]
files.sort(key=sort_numerically)

# Función para preprocesar la imagen
def preprocess_image(image_path):
    # Leer la imagen usando OpenCV
    img = cv2.imread(image_path, cv2.IMREAD_COLOR)
    
    # Convertir a escala de grises
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    
    # Aplicar binarización (thresholding)
    _, binary = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    
    # Eliminar ruido
    denoised = cv2.fastNlMeansDenoising(binary, None, 30, 7, 21)
    
    # Guardar la imagen preprocesada temporalmente
    temp_filename = "temp_image.png"
    cv2.imwrite(temp_filename, denoised)
    
    return temp_filename

# Inicializar el lector de EasyOCR
reader = easyocr.Reader(['en', 'ko'])  # Ajustar al idioma necesario

# Recorrer cada archivo en la carpeta ordenada
for filename in files:
    # Ruta completa al archivo de imagen
    image_path = os.path.join(folder_path, filename)

    # Preprocesar la imagen
    preprocessed_image_path = preprocess_image(image_path)

    # Usar EasyOCR para extraer texto
    result = reader.readtext(preprocessed_image_path, detail=0)
    joined_text = " ".join(result)

    # Abrir la imagen preprocesada
    image = Image.open(preprocessed_image_path)



    # Añadir el texto al documento
    doc.add_heading(filename, level=1)  # Añadir el nombre del archivo como encabezado
    doc.add_paragraph(joined_text)  # Añadir el texto extraído y unido

# Eliminar la imagen temporal
os.remove(preprocessed_image_path)

# Especificar la ruta completa donde se guardará el documento
output_path = 'C:/Users/hp/Desktop/TEXTO.docx'

# Guardar el documento
doc.save(output_path)