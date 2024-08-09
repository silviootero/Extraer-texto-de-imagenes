import cv2
import numpy as np
from PIL import Image
import pytesseract
from docx import Document
import os
import re

# Especificar la ruta completa al ejecutable de Tesseract
pytesseract.pytesseract.tesseract_cmd = r'D:/TESSERACT/tesseract.exe'

# Ruta a la carpeta que contiene las imágenes
folder_path = 'C:/Users/hp/Desktop/CAPS_TRADUS'

# Crear un nuevo documento de Word
doc = Document()

# Función para ordenar archivos numéricamente
def sort_numerically(value):
    parts = value.split('.')
    num = ''.join(filter(str.isdigit, parts[0]))
    return int(num)

# Obtener y ordenar los archivos numéricamente
files = [f for f in os.listdir(folder_path) if f.endswith(('.png', '.jpg', '.jpeg'))]
files.sort(key=sort_numerically)

# Función para preprocesar la imagen
def preprocess_image(image_path):
    # Leer la imagen usando OpenCV
    img = cv2.imread(image_path, cv2.IMREAD_COLOR)
    
    # Convertir a escala de grises
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    
    # Aplicar binarización (thresholding)
    _, binary = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    
    # Eliminar ruido
    denoised = cv2.fastNlMeansDenoising(binary, None, 30, 7, 21)
    
    # Guardar la imagen preprocesada temporalmente
    temp_filename = "temp_image.png"
    cv2.imwrite(temp_filename, denoised)
    
    return temp_filename

# Variable para almacenar texto acumulado
accumulated_text = ""

# Recorrer cada archivo en la carpeta ordenada
for filename in files:
    # Ruta completa al archivo de imagen
    image_path = os.path.join(folder_path, filename)

    # Preprocesar la imagen
    preprocessed_image_path = preprocess_image(image_path)

    # Abrir la imagen preprocesada
    image = Image.open(preprocessed_image_path)

    # Usar pytesseract para extraer texto
    text = pytesseract.image_to_string(image, lang='eng+kor')

    # Eliminar saltos de línea al final del texto para combinar correctamente
    text = text.strip()

    # Comprobar si hay texto acumulado
    if accumulated_text:
        # Si el nuevo texto no comienza con un espacio, lo añadimos al texto acumulado
        if not text.startswith(" "):
            accumulated_text += " " + text
        else:
            # Añadir el texto acumulado como un párrafo al documento
            doc.add_heading(filename, level=1)  # Añadir el nombre del archivo como encabezado
            doc.add_paragraph(accumulated_text)
            accumulated_text = text  # Resetear texto acumulado
    else:
        accumulated_text = text

# Añadir el último bloque de texto acumulado, si existe
if accumulated_text:
    doc.add_paragraph(accumulated_text)

# Eliminar la imagen temporal
os.remove(preprocessed_image_path)

# Guardar el documento
doc.save('TEXTO.docx')


# Ruta al archivo original
input_file = 'C:/Users/hp/Desktop/TEXTO.docx'

# Ruta al archivo modificado
output_file = 'C:/Users/hp/Desktop/TEXTO_modificado.docx'

# Cargar el documento
doc = Document(input_file)

# Crear un nuevo documento para el archivo modificado
new_doc = Document()

# Procesar cada párrafo en el documento original
for para in doc.paragraphs:
    # Reemplazar secuencias de espacios con un solo espacio (puedes ajustar esto según tus necesidades)
    modified_text = re.sub(r'\s+', ' ', para.text)
    
    # Añadir el párrafo modificado al nuevo documento
    new_doc.add_paragraph(modified_text)

# Guardar el documento modificado
new_doc.save(output_file)

print("Archivo modificado guardado en:", output_file)